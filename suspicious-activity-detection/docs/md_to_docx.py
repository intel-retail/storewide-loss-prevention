#!/usr/bin/env python3
"""Convert a Markdown file to a .docx Word document.

Scoped converter for docs/vlm-recall.md: supports ATX headings, GitHub-style
pipe tables, fenced code blocks (rendered as monospaced shaded blocks, including
mermaid), ordered/unordered lists (with nesting by indent), blockquotes, and
inline bold/inline-code. Not a full CommonMark implementation.
"""
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt as _Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MONO_FONT = "Consolas"
CODE_SHADE = "F2F2F2"


def shade_paragraph(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(paragraph, text):
    """Add text to a paragraph, honoring **bold** and `code` spans."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
            run.font.size = _Pt(9.5)
        else:
            paragraph.add_run(part)


def split_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|[\s:|-]*$", line)) and "-" in line


def convert(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = _Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # Fenced code block
        m = re.match(r"^```(.*)$", line)
        if m:
            lang = m.group(1).strip()
            i += 1
            code_lines = []
            while i < n and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if lang:
                cap = doc.add_paragraph()
                run = cap.add_run(f"[{lang}]")
                run.italic = True
                run.font.size = _Pt(8)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            for cl in code_lines:
                p = doc.add_paragraph()
                shade_paragraph(p, CODE_SHADE)
                p.paragraph_format.space_after = _Pt(0)
                p.paragraph_format.space_before = _Pt(0)
                run = p.add_run(cl if cl else " ")
                run.font.name = MONO_FONT
                run.font.size = _Pt(9)
            doc.add_paragraph()
            continue

        # Table
        if line.strip().startswith("|") and i + 1 < n and is_separator(lines[i + 1]):
            header = split_table_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            ncol = len(header)
            table = doc.add_table(rows=1, cols=ncol)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            hdr = table.rows[0].cells
            for c, htext in enumerate(header):
                hdr[c].paragraphs[0].text = ""
                add_inline(hdr[c].paragraphs[0], htext)
                for run in hdr[c].paragraphs[0].runs:
                    run.bold = True
                shade_cell(hdr[c], "D9E2F3")
            for row in rows:
                cells = table.add_row().cells
                for c in range(ncol):
                    val = row[c] if c < len(row) else ""
                    cells[c].paragraphs[0].text = ""
                    add_inline(cells[c].paragraphs[0], val)
            doc.add_paragraph()
            continue

        # Headings
        hm = re.match(r"^(#{1,6})\s+(.*)$", line)
        if hm:
            level = len(hm.group(1))
            doc.add_heading(hm.group(2).strip(), level=min(level, 4))
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            text = re.sub(r"^>\s?", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = _Pt(18)
            shade_paragraph(p, "FFF6D6")
            add_inline(p, text)
            i += 1
            continue

        # Ordered list
        om = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if om:
            indent = len(om.group(1))
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = _Pt(18 + indent * 6)
            add_inline(p, om.group(3))
            i += 1
            continue

        # Unordered list
        um = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if um:
            indent = len(um.group(1))
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = _Pt(18 + indent * 6)
            add_inline(p, um.group(2))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Paragraph (accumulate wrapped lines)
        para_lines = [line]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#{1,6}\s|```|>|\s*[-*]\s|\s*\d+\.\s|\|)", lines[i]
        ):
            para_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(s.strip() for s in para_lines))

    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "vlm-recall.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else "vlm-recall.docx"
    convert(src, dst)
