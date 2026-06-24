"""Generate a concise arch-forum brief (docx) for VLM-Recall Search with VSS."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1E, 0x6B, 0x3A)
GREY = RGBColor(0x55, 0x55, 0x55)


def h(doc, text, size=13, color=ACCENT, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def kv(doc, key, val):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(key + ": ")
    r.bold = True
    p.add_run(val)
    return p


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("VLM-Recall Search with VSS")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
sub = doc.add_paragraph()
r = sub.add_run("Architecture Forum Brief")
r.font.size = Pt(11)
r.font.color.rgb = GREY
meta = doc.add_paragraph()
r = meta.add_run("Storewide Loss Prevention \u00b7 Suspicious Activity Detection \u00b7 2026-06-24")
r.font.size = Pt(9)
r.font.color.rgb = GREY

# 1. Problem
h(doc, "1. Problem")
doc.add_paragraph(
    "Investigators need to recall store-camera footage by natural-language query, "
    "scoped to a camera/area and an absolute time window (\u201cshow aisle-7 between "
    "2\u20133pm yesterday\u201d). We want this without building and maintaining a new "
    "video index, vector DB, or mapping database of our own."
)

# 2. Approach
h(doc, "2. Approach (one line)")
doc.add_paragraph(
    "A thin, stateless ingest bridge segments and tags RTSP into short clips and "
    "uploads them to the existing Intel VSS (Video Search & Summarization) stack. "
    "VSS owns embedding, storage, similarity search, time + tag filtering, and "
    "playback URLs. No mapping DB, no duplicated index on our side."
)

# 3. Architecture
h(doc, "3. Architecture")
bullet(doc, "segments RTSP \u2192 short MP4 clips; tags each clip (camera, area, store, date-bucket) from cameras.yaml; uploads to VSS near-real-time.", bold_lead="Ingest plane (vss-recall-bridge): ")
bullet(doc, "pipeline-manager API (/manager/*), vdms-dataprep + CLIP embedding, search-ms similarity, VDMS vector DB, MinIO/Postgres storage.", bold_lead="VSS Search Stack (existing): ")
bullet(doc, "query + tags + absolute timeFilter \u2192 ranked segments with start/end + playback URL. Gradio dashboard can call VSS directly; a thin proxy is optional (auth/audit/window-padding).", bold_lead="Search path: ")
doc.add_paragraph(
    "See companion diagram: vlm-recall-hld.drawio."
).runs[0].font.color.rgb = GREY

# 4. The one dependency
h(doc, "4. The one dependency \u2014 VSS fix R1", color=RGBColor(0xB0, 0x00, 0x20))
doc.add_paragraph(
    "VSS must honor an absolute timeFilter:{start,end} on search, combinable with "
    "tags. Today normalizeTimeFilter() in the pipeline-manager drops absolute "
    "start/end unless a relative {value,unit} is supplied, so time-scoped search "
    "does not work. R1 is the only required change; the rest of this design assumes "
    "it lands."
)

# 5. Capture time without a new field
h(doc, "5. Capture time \u2014 no new ingest field needed")
doc.add_paragraph(
    "We deliberately do NOT add a capture-timestamp field at ingest (previously "
    "\u201cR2\u201d). The bridge waits for each ~5-min chunk to close, then uploads "
    "near-real-time, so VSS created_at (upload time) \u2248 capture time, bounded "
    "within one chunk. The query side pads the requested window by one chunk "
    "(WINDOW_PAD_SECONDS) so edge clips are still caught."
)
kv(doc, "R2 (caller-supplied capture timestamp)", "OPTIONAL \u2014 only for bulk backfill of old footage, where upload time \u2260 capture time.")

# 6. Why this is the right call
h(doc, "6. Why VSS-only / stateless")
bullet(doc, "no second index or mapping DB to keep in sync (no orphan/consistency risk).")
bullet(doc, "filtering, ranking, and playback URLs are produced server-side by VSS.")
bullet(doc, "smallest possible footprint we own: segment + tag + upload + proxy.")

# 7. Asks / decisions
h(doc, "7. Decisions requested")
bullet(doc, "Approve VSS R1 (absolute timeFilter) as a required upstream change.")
bullet(doc, "Confirm stateless / no-mapping-DB direction (capture time via upload + padding).")
bullet(doc, "Agree R2 is deferred and scoped to backfill only.")

# 8. How VSS search works internally
h(doc, "8. How VSS search works internally (no VLM in the query path)")
doc.add_paragraph(
    "VSS search is embedding-based vector similarity \u2014 there is no LLM/VLM call "
    "when a query runs. The VLM is used only in the separate summarization "
    "(caption) pipeline, not in search."
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Ingest (dataprep): ")
r.bold = True
p.add_run(
    "VSS samples frames from each clip, embeds them with a CLIP-style multimodal "
    "model (multimodal-embedding-serving), and stores frame vectors + metadata "
    "(video_id, timestamp, created_at, tags) in the VDMS vector DB."
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Query (search-ms): ")
r.bold = True
p.add_run("runs the following steps \u2014")
bullet(doc, "the query text is embedded with the SAME CLIP model, so text and image share one vector space.", bold_lead="1. Embed query: ")
bullet(doc, "VDMS runs approximate-nearest-neighbour over frame vectors, with timeFilter:{start,end} + tags applied as a metadata pre-filter (this is exactly where R1 must take effect).", bold_lead="2. ANN + filter: ")
bullet(doc, "matching frames are grouped into ~8s temporal segments and re-scored, returning ranked segments with segment_start/end + score + playback URL.", bold_lead="3. Aggregate: ")
doc.add_paragraph(
    "Net: CLIP text\u2192image matching + metadata filtering + temporal aggregation. "
    "Great for visual concepts (\u201cperson in red jacket\u201d, \u201cforklift\u201d, "
    "\u201cspill\u201d); it does not reason or count \u2014 that remains a VLM task outside "
    "the search path."
).runs[0].font.size = Pt(10)

# Footer note
n = doc.add_paragraph()
n.paragraph_format.space_before = Pt(12)
rn = n.add_run("References: vlm-recall.md (rationale), vlm-recall-build-spec.md (implementation spec), vlm-recall-hld.drawio (diagram).")
rn.italic = True
rn.font.size = Pt(8.5)
rn.font.color.rgb = GREY

out = "vlm-recall-arch-forum.docx"
doc.save(out)
print("wrote", out)
